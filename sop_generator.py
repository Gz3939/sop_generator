import streamlit as st
import os
import json
import pickle
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
from io import BytesIO
import tempfile
from datetime import datetime

class SOPGenerator:
    def __init__(self):
        self.steps = []
        self.autosave_file = "autosave_data.pkl"
        self.config_file = "sop_config.json"
        
    def save_to_file(self, sop_data):
        """自動保存SOP數據到本地文件"""
        try:
            # 保存步驟數據（包含圖片）
            with open(self.autosave_file, 'wb') as f:
                pickle.dump(self.steps, f)
            
            # 保存配置數據（不包含圖片）
            config_data = {
                'sop_title': sop_data.get('sop_title', ''),
                'author': sop_data.get('author', ''),
                'creation_date': sop_data.get('creation_date', '').strftime('%Y-%m-%d') if sop_data.get('creation_date') else '',
                'purpose': sop_data.get('purpose', ''),
                'scope': sop_data.get('scope', ''),
                'last_saved': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(config_data, f, ensure_ascii=False, indent=2)
                
            return True
        except Exception as e:
            st.error(f"自動保存失敗: {str(e)}")
            return False
    
    def load_from_file(self):
        """從本地文件加載SOP數據"""
        try:
            # 加載步驟數據
            if os.path.exists(self.autosave_file):
                with open(self.autosave_file, 'rb') as f:
                    self.steps = pickle.load(f)
            
            # 加載配置數據
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    config_data = json.load(f)
                return config_data
            
            return None
        except Exception as e:
            st.error(f"讀取保存數據失敗: {str(e)}")
            return None
    
    def clear_autosave(self):
        """清除自動保存文件"""
        try:
            if os.path.exists(self.autosave_file):
                os.remove(self.autosave_file)
            if os.path.exists(self.config_file):
                os.remove(self.config_file)
            return True
        except:
            return False
        
    def add_step(self, image, description, step_number):
        """添加一個步驟到SOP"""
        self.steps.append({
            'step_number': step_number,
            'image': image,
            'description': description
        })
    
    def generate_word_document(self, title="標準作業程序 (SOP)"):
        """生成Word文檔"""
        doc = Document()
        
        # 添加標題
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加日期
        date_paragraph = doc.add_paragraph(f"建立日期: {st.session_state.get('creation_date', '未設定')}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # 添加作者
        author_paragraph = doc.add_paragraph(f"建立者: {st.session_state.get('author', '未設定')}")
        author_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # 添加目的
        if st.session_state.get('purpose'):
            doc.add_heading('目的', level=1)
            doc.add_paragraph(st.session_state.get('purpose'))
        
        # 添加適用範圍
        if st.session_state.get('scope'):
            doc.add_heading('適用範圍', level=1)
            doc.add_paragraph(st.session_state.get('scope'))
        
        # 添加步驟
        doc.add_heading('操作步驟', level=1)
        
        # 創建臨時目錄用於存放圖片
        temp_dir = "temp_images"
        try:
            if not os.path.exists(temp_dir):
                os.makedirs(temp_dir)
        except:
            temp_dir = "."  # 如果無法創建目錄，使用當前目錄
        
        temp_files = []  # 記錄需要清理的臨時文件
        
        for step in self.steps:
            try:
                # 步驟標題
                step_heading = doc.add_heading(f"步驟 {step['step_number']}", level=2)
                
                # 添加圖片
                if step['image'] is not None:
                    # 生成唯一的臨時文件名
                    import time
                    timestamp = str(int(time.time() * 1000))
                    temp_filename = os.path.join(temp_dir, f"step_{step['step_number']}_{timestamp}.png")
                    
                    try:
                        # 保存圖片
                        step['image'].save(temp_filename, format='PNG')
                        temp_files.append(temp_filename)
                        
                        # 計算圖片大小（最大寬度6英寸）
                        img_width = min(6, step['image'].width / 100)
                        
                        # 添加圖片到文檔
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                        run.add_picture(temp_filename, width=Inches(img_width))
                        
                    except Exception as img_error:
                        # 如果圖片處理失敗，添加錯誤說明
                        error_paragraph = doc.add_paragraph(f"[圖片載入錯誤: {str(img_error)}]")
                        error_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 添加步驟描述
                if step['description']:
                    desc_paragraph = doc.add_paragraph(step['description'])
                    desc_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # 添加空行
                doc.add_paragraph()
                
            except Exception as step_error:
                # 如果步驟處理失敗，添加錯誤說明並繼續
                error_heading = doc.add_heading(f"步驟 {step['step_number']} (處理錯誤)", level=2)
                error_paragraph = doc.add_paragraph(f"步驟處理錯誤: {str(step_error)}")
                if step.get('description'):
                    doc.add_paragraph(step['description'])
                doc.add_paragraph()
        
        # 清理所有臨時文件
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except:
                pass  # 忽略清理錯誤
        
        # 清理臨時目錄（如果是我們創建的且為空）
        try:
            if temp_dir != "." and os.path.exists(temp_dir):
                os.rmdir(temp_dir)
        except:
            pass  # 忽略清理錯誤
        
        return doc

def main():
    st.set_page_config(
        page_title="SOP生成系統",
        page_icon="📋",
        layout="wide"
    )
    
    st.title("📋 SOP生成系統")
    st.markdown("---")
    
    # 初始化session state
    if 'sop_generator' not in st.session_state:
        st.session_state.sop_generator = SOPGenerator()
        
        # 嘗試加載之前保存的數據
        saved_config = st.session_state.sop_generator.load_from_file()
        if saved_config:
            st.session_state.sop_title = saved_config.get('sop_title', '標準作業程序')
            st.session_state.author = saved_config.get('author', '')
            # 處理日期
            if saved_config.get('creation_date'):
                try:
                    st.session_state.creation_date = datetime.strptime(saved_config['creation_date'], '%Y-%m-%d').date()
                except:
                    st.session_state.creation_date = datetime.now().date()
            else:
                st.session_state.creation_date = datetime.now().date()
            st.session_state.purpose = saved_config.get('purpose', '')
            st.session_state.scope = saved_config.get('scope', '')
            
            # 顯示恢復提示
            if st.session_state.sop_generator.steps:
                st.success(f"✅ 已恢復之前的工作進度！(上次保存: {saved_config.get('last_saved', '未知')})")
    
    if 'current_step' not in st.session_state:
        st.session_state.current_step = 1
    
    # 初始化其他session state變數
    if 'sop_title' not in st.session_state:
        st.session_state.sop_title = '標準作業程序'
    if 'author' not in st.session_state:
        st.session_state.author = ''
    if 'creation_date' not in st.session_state:
        st.session_state.creation_date = datetime.now().date()
    if 'purpose' not in st.session_state:
        st.session_state.purpose = ''
    if 'scope' not in st.session_state:
        st.session_state.scope = ''
    
    # 側邊欄 - SOP基本資訊
    with st.sidebar:
        st.header("📝 SOP基本資訊")
        
        st.session_state.sop_title = st.text_input(
            "SOP標題", 
            value=st.session_state.get('sop_title', '標準作業程序')
        )
        
        st.session_state.author = st.text_input(
            "建立者", 
            value=st.session_state.get('author', '')
        )
        
        st.session_state.creation_date = st.date_input(
            "建立日期",
            value=st.session_state.get('creation_date', datetime.now().date())
        )
        
        st.session_state.purpose = st.text_area(
            "目的",
            value=st.session_state.get('purpose', ''),
            height=100
        )
        
        st.session_state.scope = st.text_area(
            "適用範圍",
            value=st.session_state.get('scope', ''),
            height=100
        )
        
        # 自動保存功能
        def auto_save():
            sop_data = {
                'sop_title': st.session_state.sop_title,
                'author': st.session_state.author,
                'creation_date': st.session_state.creation_date,
                'purpose': st.session_state.purpose,
                'scope': st.session_state.scope
            }
            st.session_state.sop_generator.save_to_file(sop_data)
        
        # 當有任何變更時自動保存
        if (st.session_state.get('prev_sop_title') != st.session_state.sop_title or
            st.session_state.get('prev_author') != st.session_state.author or
            st.session_state.get('prev_creation_date') != st.session_state.creation_date or
            st.session_state.get('prev_purpose') != st.session_state.purpose or
            st.session_state.get('prev_scope') != st.session_state.scope):
            
            auto_save()
            
            # 更新previous值
            st.session_state.prev_sop_title = st.session_state.sop_title
            st.session_state.prev_author = st.session_state.author
            st.session_state.prev_creation_date = st.session_state.creation_date
            st.session_state.prev_purpose = st.session_state.purpose
            st.session_state.prev_scope = st.session_state.scope
        
        st.markdown("---")
        
        # 自動保存狀態
        col1, col2 = st.columns(2)
        with col1:
            if st.button("💾 手動保存", help="立即保存當前進度"):
                auto_save()
                st.success("保存成功！")
        
        with col2:
            if st.button("🗑️ 清除保存", help="清除所有保存的數據"):
                if st.session_state.sop_generator.clear_autosave():
                    st.session_state.sop_generator.steps = []
                    st.session_state.current_step = 1
                    st.success("已清除所有數據")
                    st.rerun()
        
        # 顯示自動保存狀態
        if os.path.exists(st.session_state.sop_generator.config_file):
            try:
                with open(st.session_state.sop_generator.config_file, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    last_saved = config.get('last_saved', '未知')
                    st.caption(f"🕒 上次保存: {last_saved}")
            except:
                pass
        
        st.markdown("---")
        
        # 顯示已添加的步驟
        st.header("📚 已添加步驟")
        if st.session_state.sop_generator.steps:
            for i, step in enumerate(st.session_state.sop_generator.steps):
                with st.expander(f"步驟 {step['step_number']}"):
                    if step['image']:
                        st.image(step['image'], width=200)
                    st.write(step['description'])
                    if st.button(f"刪除步驟 {step['step_number']}", key=f"delete_{i}"):
                        st.session_state.sop_generator.steps.pop(i)
                        
                        # 自動保存
                        sop_data = {
                            'sop_title': st.session_state.sop_title,
                            'author': st.session_state.author,
                            'creation_date': st.session_state.creation_date,
                            'purpose': st.session_state.purpose,
                            'scope': st.session_state.scope
                        }
                        st.session_state.sop_generator.save_to_file(sop_data)
                        
                        st.rerun()
        else:
            st.info("尚未添加任何步驟")
    
    # 主要內容區域
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📸 添加新步驟")
        
        # 步驟編號
        step_number = st.number_input(
            "步驟編號", 
            min_value=1, 
            value=st.session_state.current_step,
            step=1
        )
        
        # 圖片上傳
        uploaded_image = st.file_uploader(
            "上傳步驟圖片",
            type=['png', 'jpg', 'jpeg', 'gif'],
            help="支援PNG、JPG、JPEG、GIF格式"
        )
        
        # 顯示上傳的圖片
        if uploaded_image is not None:
            image = Image.open(uploaded_image)
            st.image(image, caption=f"步驟 {step_number} 圖片預覽", use_column_width=True)
        
        # 步驟描述
        step_description = st.text_area(
            "步驟說明",
            height=150,
            placeholder="請輸入這個步驟的詳細說明..."
        )
        
        # 添加步驟按鈕
        if st.button("➕ 添加步驟", type="primary"):
            if uploaded_image is not None and step_description.strip():
                image = Image.open(uploaded_image)
                st.session_state.sop_generator.add_step(
                    image, 
                    step_description, 
                    step_number
                )
                st.session_state.current_step = step_number + 1
                
                # 自動保存
                sop_data = {
                    'sop_title': st.session_state.sop_title,
                    'author': st.session_state.author,
                    'creation_date': st.session_state.creation_date,
                    'purpose': st.session_state.purpose,
                    'scope': st.session_state.scope
                }
                st.session_state.sop_generator.save_to_file(sop_data)
                
                st.success(f"步驟 {step_number} 已成功添加並保存！")
                st.rerun()
            else:
                st.error("請上傳圖片並填寫步驟說明")
    
    with col2:
        st.header("📄 預覽與生成")
        
        if st.session_state.sop_generator.steps:
            st.subheader("SOP預覽")
            
            # 顯示SOP標題
            st.markdown(f"### {st.session_state.sop_title}")
            st.markdown(f"**建立者:** {st.session_state.author}")
            st.markdown(f"**建立日期:** {st.session_state.creation_date}")
            
            if st.session_state.purpose:
                st.markdown(f"**目的:** {st.session_state.purpose}")
            
            if st.session_state.scope:
                st.markdown(f"**適用範圍:** {st.session_state.scope}")
            
            st.markdown("---")
            st.markdown("### 操作步驟")
            
            # 顯示所有步驟
            for step in sorted(st.session_state.sop_generator.steps, key=lambda x: x['step_number']):
                with st.container():
                    st.markdown(f"#### 步驟 {step['step_number']}")
                    if step['image']:
                        st.image(step['image'], width=300)
                    st.markdown(step['description'])
                    st.markdown("---")
            
            # 生成Word文檔按鈕
            if st.button("📄 生成Word文檔", type="primary"):
                try:
                    doc = st.session_state.sop_generator.generate_word_document(st.session_state.sop_title)
                    
                    # 保存文檔到內存
                    doc_buffer = BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    
                    # 提供下載按鈕
                    st.download_button(
                        label="⬇️ 下載Word文檔",
                        data=doc_buffer.getvalue(),
                        file_name=f"{st.session_state.sop_title}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    st.success("Word文檔生成成功！點擊上方按鈕下載。")
                    
                except Exception as e:
                    st.error(f"生成Word文檔時發生錯誤: {str(e)}")
        
        else:
            st.info("請先添加至少一個步驟才能生成SOP")
    
    # 清除所有步驟按鈕
    st.markdown("---")
    if st.session_state.sop_generator.steps:
        if st.button("🗑️ 清除所有步驟", type="secondary"):
            st.session_state.sop_generator.steps = []
            st.session_state.current_step = 1
            st.success("所有步驟已清除")
            st.rerun()

if __name__ == "__main__":
    main()
